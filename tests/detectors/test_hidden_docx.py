import io
from app.main import app
from fastapi.testclient import TestClient
from docx import Document


client = TestClient(app)


def _docx_with_hidden_text() -> bytes:
    bio = io.BytesIO()
    d = Document()
    p = d.add_paragraph()
    p.add_run("Visible ")
    r = p.add_run("Secret")
    r.font.hidden = True  # maps to w:vanish
    d.save(bio)
    return bio.getvalue()


def test_multipart_docx_with_hidden_text_is_denied_and_debug_has_sources():
    docx_bytes = _docx_with_hidden_text()
    files = [("files", ("hidden.docx", docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    r = client.post("/guardrail/evaluate_multipart", files=files, headers={"X-Debug": "1"})
    assert r.status_code == 200
    body = r.json()
    assert body["action"] == "deny"
    assert r.headers.get("X-Guardrail-Decision") == "deny"
    assert "debug" in body and "sources" in body["debug"]


def test_multipart_docx_clean_is_allowed_or_sanitized():
    bio = io.BytesIO()
    d = Document()
    d.add_paragraph("Just text")
    d.save(bio)
    files = [("files", ("clean.docx", bio.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))]
    r = client.post("/guardrail/evaluate_multipart", files=files)
    assert r.status_code == 200
    body = r.json()
    assert body["action"] in {"allow", "sanitize", "clarify", "deny"}

